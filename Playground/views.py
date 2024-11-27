from flask import Blueprint, current_app, render_template, redirect, url_for, request, session, jsonify
from flask_login import login_required, current_user
from .util import convert_file_to_thumbnail, parse_page_ranges, extract_text, generate_questions
from .exam_util import exam_convert_file_to_thumbnail, exam_parse_page_ranges, exam_extract_text, exam_generate_questions
import os
from flask import send_file
from docx import Document

views = Blueprint('views', __name__)


@views.route('/')
def index():
    return render_template('index.html')

@views.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html', user=current_user)

@views.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'GET':
        return render_template('upload.html') 
    
    elif request.method == 'POST':
        if 'input_file' not in request.files:
            return redirect(request.url)

        file = request.files['input_file']
        if file.filename == '':
            return redirect(request.url)
        
        session['file_path'] = os.path.join(current_app.config['UPLOAD_FOLDER'], file.filename)
        file.save(session['file_path'])

        return redirect(url_for('views.selection', file_name=file.filename))


@views.route('/selection',  methods=['GET', 'POST'])
def selection():
    if request.method == 'GET':
        filename = request.args.get('file_name')
        thumbnails = convert_file_to_thumbnail(session['file_path'], current_app.config['THUMBNAIL_FOLDER'], start_page=0, end_page=10)

        return render_template('preview.html', filename=filename, thumbnails=thumbnails)
    
    elif request.method == 'POST':
        filename = request.form.get('filename')
        page_selection = request.form.get('page-selection')
        pages = request.form.get('pages')
        pages = parse_page_ranges(pages)

        question_types = request.form.getlist('ques-type')
        question_quantities = request.form.getlist('ques-num')
         # Process questions and quantities
        questions = [{'type': qt, 'quantity': int(qn)} for qt, qn in zip(question_types, question_quantities) if qn.isdigit()]
        
        # NOTE: FORM VALIDATION HERE

        print(f"Filename: {filename}")
        print(f"Page Selection: {page_selection}")
        print(f"Pages: {pages}")
        print("Questions:", questions)

        text = ''
        try:
            text = extract_text(session['file_path'], pages)
        except KeyError:
            return jsonify({'message': 'Return to Upload Page'}), 400


        session['questions'] = questions
        session['text'] = text

        return redirect(url_for('views.download'))
        #return jsonify({'questions': questions, 'text': text})


@views.route('/download')
def download():
    questions = session.get('questions', [])
    text = session.get('text', '')

    # For debugging purpose, making sure correct values are being passed
    for question in questions:
        question_type = question.get('type')
        num_questions = question.get('quantity')
        print(f"Preparing to generate {num_questions} {question_type} questions.")

    generated_questions = generate_questions(questions, text)

    #return render_template('download.html', generated_questions=generated_questions) # download page shows the generated question and button to continue to quiz
    return render_template('generated.html', generated_questions=generated_questions)


@views.route('/done')
def done():
    return render_template('done.html') 


@views.route('/review-questions')
def review_questions():
    ques_type = request.args.get('ques_type')
    return render_template('review-ques.html', ques_type=ques_type)



@views.route('/quiz-complete', methods=['GET', 'POST'])
def quiz_complete():
    score = request.args.get('score')
    total = request.args.get('total')

    return render_template('quiz-complete.html', score=score, total=total)




# ROUTES FOR FETCHING DATA 
@views.route('/selection/<int:page>')
def load_thumbnails(page):
    end = page + 10
    thumbnails = convert_file_to_thumbnail(session['file_path'], current_app.config['THUMBNAIL_FOLDER'], start_page=page, end_page=end)
    return jsonify(thumbnails=thumbnails)



@views.route('/quiz-complete/responses', methods=['POST'])
def responses():
    result = request.get_json()
    score = result.get('score')
    total = result.get('totalQuestion')

    return redirect( url_for('views.quiz_complete', score=score, total=total) )








#EXAM ROUTES
@views.route('/exam_upload', methods=['GET', 'POST'])
def exam_upload():
    if request.method == 'GET':
        return render_template('exam_upload.html') 
    
    elif request.method == 'POST':
        if 'input_file' not in request.files:
            return redirect(request.url)

        file = request.files['input_file']
        if file.filename == '':
            return redirect(request.url)
        
        session['file_path'] = os.path.join(current_app.config['UPLOAD_FOLDER'], file.filename)
        file.save(session['file_path'])

        return redirect(url_for('views.exam_selection', file_name=file.filename))


@views.route('/exam_selection',  methods=['GET', 'POST'])
def exam_selection():
    if request.method == 'GET':
        filename = request.args.get('file_name')
        thumbnails = exam_convert_file_to_thumbnail(session['file_path'], current_app.config['THUMBNAIL_FOLDER'], start_page=0, end_page=10)

        return render_template('exam_preview.html', filename=filename, thumbnails=thumbnails)
    
    elif request.method == 'POST':
        filename = request.form.get('filename')
        page_selection = request.form.get('page-selection')
        pages = request.form.get('pages')
        pages = exam_parse_page_ranges(pages)

        question_types = request.form.getlist('ques-type')
        question_quantities = request.form.getlist('ques-num')
         # Process questions and quantities
        questions = [{'type': qt, 'quantity': int(qn)} for qt, qn in zip(question_types, question_quantities) if qn.isdigit()]
        
        # NOTE: FORM VALIDATION HERE

        print(f"Filename: {filename}")
        print(f"Page Selection: {page_selection}")
        print(f"Pages: {pages}")
        print("Questions:", questions)

        text = ''
        try:
            text = exam_extract_text(session['file_path'], pages)
        except KeyError:
            return jsonify({'message': 'Return to Upload Page'}), 400


        session['questions'] = questions
        session['text'] = text

        return redirect(url_for('views.exam_download'))
        #return jsonify({'questions': questions, 'text': text})


@views.route('/exam_download')
def exam_download():
    questions = session.get('questions', [])
    text = session.get('text', '')

    # For debugging purpose, making sure correct values are being passed
    for question in questions:
        question_type = question.get('type')
        num_questions = question.get('quantity')
        print(f"Preparing to generate {num_questions} {question_type} questions.")

    # Generate the questions using the function
    generated_questions = exam_generate_questions(questions, text)

    # For debugging, check the structure of the questions being passed to the template
    print(f"Generated questions: {generated_questions}")
    session['questions'] = generated_questions

    # Pass the generated questions to the template
    return render_template('exam_generated.html', generated_questions=generated_questions)


@views.route('/download_document', methods=['GET'])
def download_document():
    questions = session.get('questions', [])  # Generated questions
    text = session.get('text', '')  # Extracted text

    # Create a new Document
    doc = Document()
    doc.add_heading('Generated Questions and Text', 0)

    # Add generated questions
    for question_set in questions:
        doc.add_heading(f'{question_set["type"]}',level=2)

        for question in question_set['questions']:
            doc.add_paragraph(f'{question["question"]}',style='ListNumber')

            # Add options if present
            if 'options' in question:
                for letter, option in question['options'].items():
                    doc.add_paragraph(f'{letter}: {option}')

            # Add the answer
            doc.add_paragraph(f'Answer: {question.get("answer", "N/A")}')

    # Save the document temporarily
    file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], 'generated_questions_and_text.docx')
    doc.save(file_path)

    # Send the document to the user for download
    return send_file(file_path, as_attachment=True)
